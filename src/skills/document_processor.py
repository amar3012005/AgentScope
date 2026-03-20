"""
Document Processing Skill for BLAIQ
Handles document chunking, embedding, and storage in Qdrant + Neo4j
"""

import os
import re
import json
import hashlib
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

# Try to import document parsers
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentChunker:
    """Intelligent document chunking with LLM assistance"""

    def __init__(self, llm_client=None, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.llm_client = llm_client

    def extract_text(self, file_path: str, file_content: bytes = None) -> str:
        """Extract text from various document formats"""
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == '.pdf' and PYPDF_AVAILABLE:
            return self._extract_pdf(file_content or file_path)
        elif suffix == '.docx' and DOCX_AVAILABLE:
            return self._extract_docx(file_content or file_path)
        elif suffix == '.txt':
            if isinstance(file_content, bytes):
                return file_content.decode('utf-8', errors='ignore')
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        elif suffix == '.md':
            if isinstance(file_content, bytes):
                return file_content.decode('utf-8', errors='ignore')
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        else:
            # Try as text
            if file_content:
                return file_content.decode('utf-8', errors='ignore')
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

    def _extract_pdf(self, source) -> str:
        """Extract text from PDF"""
        if isinstance(source, bytes):
            from io import BytesIO
            pdf_file = BytesIO(source)
        else:
            pdf_file = source

        reader = pypdf.PdfReader(pdf_file)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)

    def _extract_docx(self, source) -> str:
        """Extract text from DOCX"""
        if isinstance(source, bytes):
            from io import BytesIO
            doc = DocxDocument(BytesIO(source))
        else:
            doc = DocxDocument(source)

        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        return "\n\n".join(text_parts)

    def create_chunks(self, text: str, doc_id: str, source_file: str,
                      use_llm: bool = True) -> List[Dict[str, Any]]:
        """Create intelligent chunks from text"""

        if use_llm and self.llm_client:
            return self._llm_assisted_chunking(text, doc_id, source_file)
        else:
            return self._semantic_chunking(text, doc_id, source_file)

    def _semantic_chunking(self, text: str, doc_id: str, source_file: str) -> List[Dict[str, Any]]:
        """Semantic chunking using paragraph and sentence boundaries"""
        chunks = []

        # Split by paragraphs first
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

        current_chunk = []
        current_length = 0
        chunk_index = 0

        for para in paragraphs:
            para_length = len(para)

            # If adding this paragraph exceeds chunk size, finalize current chunk
            if current_length + para_length > self.chunk_size and current_chunk:
                chunk_text = "\n\n".join(current_chunk)
                chunks.append(self._create_chunk_dict(
                    chunk_text, doc_id, source_file, chunk_index
                ))

                # Keep overlap
                current_chunk = current_chunk[-2:] if len(current_chunk) > 2 else current_chunk
                current_length = sum(len(c) for c in current_chunk)
                chunk_index += 1

            current_chunk.append(para)
            current_length += para_length

        # Don't forget the last chunk
        if current_chunk:
            chunk_text = "\n\n".join(current_chunk)
            chunks.append(self._create_chunk_dict(
                chunk_text, doc_id, source_file, chunk_index
            ))

        return chunks

    def _llm_assisted_chunking(self, text: str, doc_id: str, source_file: str) -> List[Dict[str, Any]]:
        """Use LLM to identify semantic boundaries for better chunks"""
        # For now, fall back to semantic chunking
        # In a full implementation, this would call an LLM to identify sections
        return self._semantic_chunking(text, doc_id, source_file)

    def _create_chunk_dict(self, text: str, doc_id: str, source_file: str,
                          index: int) -> Dict[str, Any]:
        """Create a chunk dictionary with metadata"""
        chunk_id = f"{doc_id}_chunk_{index:03d}"
        return {
            "chunk_id": chunk_id,
            "doc_id": doc_id,
            "chunk_index": index,
            "text": text,
            "metadata": {
                "source_file": source_file,
                "chunk_size": len(text),
                "char_count": len(text),
                "word_count": len(text.split()),
            }
        }


class DocumentProcessor:
    """Main document processing interface"""

    def __init__(self, qdrant_client=None, neo4j_driver=None,
                 embedding_model=None, llm_client=None):
        self.qdrant_client = qdrant_client
        self.neo4j_driver = neo4j_driver
        self.embedding_model = embedding_model
        self.chunker = DocumentChunker(llm_client=llm_client)

    def process_document(self, file_path: str, file_content: bytes = None,
                        tenant_id: str = "default", metadata: Dict = None) -> Dict[str, Any]:
        """Process a document end-to-end"""

        # Generate document ID
        doc_id = self._generate_doc_id(file_path, tenant_id)

        # Extract text
        logger.info(f"Extracting text from {file_path}")
        text = self.chunker.extract_text(file_path, file_content)

        if not text.strip():
            return {
                "status": "error",
                "message": "No text extracted from document",
                "doc_id": doc_id
            }

        # Create chunks
        logger.info(f"Creating chunks for {file_path}")
        chunks = self.chunker.create_chunks(text, doc_id, file_path)

        # Generate embeddings and store in Qdrant
        if self.qdrant_client:
            logger.info(f"Storing {len(chunks)} chunks in Qdrant")
            self._store_in_qdrant(chunks, tenant_id)

        # Extract entities and store in Neo4j
        if self.neo4j_driver:
            logger.info("Extracting entities and storing in Neo4j")
            self._store_in_neo4j(chunks, tenant_id, doc_id, file_path)

        return {
            "status": "success",
            "doc_id": doc_id,
            "chunks_created": len(chunks),
            "total_chars": len(text),
            "tenant_id": tenant_id,
            "source_file": file_path
        }

    def _generate_doc_id(self, file_path: str, tenant_id: str) -> str:
        """Generate unique document ID"""
        content = f"{tenant_id}:{file_path}:{os.path.getmtime(file_path) if os.path.exists(file_path) else ''}"
        return hashlib.sha256(content.encode()).hexdigest()[:16]

    def _store_in_qdrant(self, chunks: List[Dict], tenant_id: str):
        """Store chunks with embeddings in Qdrant"""
        from qdrant_client.models import PointStruct

        points = []
        for chunk in chunks:
            # Generate embedding
            embedding = self._get_embedding(chunk["text"])

            point = PointStruct(
                id=chunk["chunk_id"],
                vector=embedding,
                payload={
                    **chunk["metadata"],
                    "text": chunk["text"],
                    "doc_id": chunk["doc_id"],
                    "chunk_index": chunk["chunk_index"],
                    "tenant_id": tenant_id,
                }
            )
            points.append(point)

        # Batch upsert
        if points:
            collection_name = os.getenv("QDRANT_COLLECTION", "graphrag_chunks")
            self.qdrant_client.upsert(collection_name=collection_name, points=points)

    def _get_embedding(self, text: str) -> List[float]:
        """Generate embedding for text"""
        if self.embedding_model:
            return self.embedding_model.embed_query(text)

        # Fallback: return zeros (should use actual embedding model)
        return [0.0] * 768

    def _store_in_neo4j(self, chunks: List[Dict], tenant_id: str,
                       doc_id: str, source_file: str):
        """Store document structure and entities in Neo4j"""
        with self.neo4j_driver.session() as session:
            # Create document node
            session.run("""
                MERGE (d:Document {id: $doc_id})
                SET d.source_file = $source_file,
                    d.tenant_id = $tenant_id,
                    d.chunk_count = $chunk_count,
                    d.created_at = datetime()
            """, doc_id=doc_id, source_file=source_file,
                 tenant_id=tenant_id, chunk_count=len(chunks))

            # Create chunk nodes and relationships
            for chunk in chunks:
                session.run("""
                    MATCH (d:Document {id: $doc_id})
                    MERGE (c:Chunk {id: $chunk_id})
                    SET c.text = $text,
                        c.chunk_index = $chunk_index,
                        c.tenant_id = $tenant_id
                    MERGE (d)-[:HAS_CHUNK]->(c)
                """, doc_id=doc_id, chunk_id=chunk["chunk_id"],
                     text=chunk["text"][:1000],  # Truncate for graph storage
                     chunk_index=chunk["chunk_index"],
                     tenant_id=tenant_id)


def process_upload(file_path: str, file_content: bytes, tenant_id: str,
                  qdrant_config: Dict, neo4j_config: Dict) -> Dict[str, Any]:
    """
    Main entry point for processing uploaded files

    Args:
        file_path: Path or filename of the uploaded file
        file_content: Raw bytes of the file
        tenant_id: Tenant identifier for multi-tenant isolation
        qdrant_config: Dict with qdrant_url, qdrant_api_key, collection_name
        neo4j_config: Dict with neo4j_uri, neo4j_user, neo4j_password

    Returns:
        Dict with processing results
    """
    from qdrant_client import QdrantClient
    from neo4j import GraphDatabase

    # Initialize clients
    qdrant_client = None
    neo4j_driver = None

    try:
        if qdrant_config.get("qdrant_url"):
            qdrant_client = QdrantClient(
                url=qdrant_config["qdrant_url"],
                api_key=qdrant_config.get("qdrant_api_key")
            )
    except Exception as e:
        logger.error(f"Failed to connect to Qdrant: {e}")

    try:
        if neo4j_config.get("neo4j_uri"):
            neo4j_driver = GraphDatabase.driver(
                neo4j_config["neo4j_uri"],
                auth=(neo4j_config.get("neo4j_user", "neo4j"),
                      neo4j_config.get("neo4j_password", ""))
            )
    except Exception as e:
        logger.error(f"Failed to connect to Neo4j: {e}")

    # Process document
    processor = DocumentProcessor(
        qdrant_client=qdrant_client,
        neo4j_driver=neo4j_driver
    )

    result = processor.process_document(
        file_path=file_path,
        file_content=file_content,
        tenant_id=tenant_id
    )

    # Cleanup
    if neo4j_driver:
        neo4j_driver.close()

    return result
